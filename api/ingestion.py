import io
import uuid
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from datetime import datetime

# Document processing imports
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.layout import LAParams
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

from api.models import FAQEntry, KBEntry
from api.storage import FileStorageManager


class DocumentParser:
    """Base class for document parsers"""
    
    def can_parse(self, filename: str) -> bool:
        """Check if this parser can handle the file"""
        raise NotImplementedError
    
    def parse(self, content: bytes, filename: str) -> str:
        """Parse document content and return text"""
        raise NotImplementedError


class PDFParser(DocumentParser):
    """PDF document parser using pdfminer.six"""
    
    def can_parse(self, filename: str) -> bool:
        return PDFMINER_AVAILABLE and filename.lower().endswith('.pdf')
    
    def parse(self, content: bytes, filename: str) -> str:
        """Extract text from PDF"""
        if not PDFMINER_AVAILABLE:
            raise RuntimeError("pdfminer.six not available for PDF parsing")
        
        try:
            # Use LAParams for better text extraction
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                char_margin=2.0,
                boxes_flow=0.5,
                all_texts=False
            )
            
            text = pdf_extract_text(io.BytesIO(content), laparams=laparams)
            return text.strip()
        
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF {filename}: {str(e)}")


class DOCXParser(DocumentParser):
    """DOCX document parser using python-docx"""
    
    def can_parse(self, filename: str) -> bool:
        return DOCX_AVAILABLE and filename.lower().endswith('.docx')
    
    def parse(self, content: bytes, filename: str) -> str:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available for DOCX parsing")
        
        try:
            doc = DocxDocument(io.BytesIO(content))
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n\n".join(paragraphs)
        
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX {filename}: {str(e)}")


class TextParser(DocumentParser):
    """Plain text and markdown parser"""
    
    def can_parse(self, filename: str) -> bool:
        return filename.lower().endswith(('.txt', '.md', '.markdown'))
    
    def parse(self, content: bytes, filename: str) -> str:
        """Parse text/markdown content"""
        try:
            text = content.decode('utf-8')
            
            # Basic markdown processing if available
            if MARKDOWN_AVAILABLE and filename.lower().endswith(('.md', '.markdown')):
                # Convert markdown to plain text (removes formatting)
                md = markdown.Markdown()
                html = md.convert(text)
                # Simple HTML tag removal
                import re
                text = re.sub(r'<[^>]+>', '', html)
            
            return text.strip()
        
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return content.decode(encoding).strip()
                except UnicodeDecodeError:
                    continue
            
            raise RuntimeError(f"Unable to decode text file {filename}")


class TextSplitter:
    """Text splitting utilities"""
    
    @staticmethod
    def split_by_tokens(text: str, max_chunk_chars: int = 1200, overlap_chars: int = 120) -> List[str]:
        """Split text into chunks by character count with overlap"""
        if len(text) <= max_chunk_chars:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chunk_chars
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Try to break at sentence boundary
            chunk = text[start:end]
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            
            # Choose the latest sentence/paragraph boundary
            break_point = max(last_period, last_newline)
            
            if break_point > start + max_chunk_chars // 2:  # Only if it's not too early
                end = start + break_point + 1
            
            chunks.append(text[start:end].strip())
            
            # Move start with overlap
            start = max(start + max_chunk_chars - overlap_chars, end - overlap_chars)
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    @staticmethod
    def split_by_headings(text: str) -> List[Tuple[str, str]]:
        """Split text by headings, return (heading, content) pairs"""
        import re
        
        # Look for markdown-style headings
        lines = text.split('\n')
        sections = []
        current_heading = "Introduction"
        current_content = []
        
        for line in lines:
            line = line.strip()
            
            # Check for markdown headings
            if re.match(r'^#{1,6}\s+', line):
                # Save previous section
                if current_content:
                    sections.append((current_heading, '\n'.join(current_content).strip()))
                
                current_heading = re.sub(r'^#{1,6}\s+', '', line)
                current_content = []
            
            # Check for underlined headings
            elif len(line) > 0 and len(current_content) == 1 and all(c in '=-' for c in line):
                current_heading = current_content[0]
                current_content = []
            
            else:
                if line:  # Skip empty lines at start of sections
                    current_content.append(line)
        
        # Add final section
        if current_content:
            sections.append((current_heading, '\n'.join(current_content).strip()))
        
        return sections if sections else [("Document", text)]


class IngestionPipeline:
    """Document ingestion pipeline"""
    
    def __init__(self, storage_manager: FileStorageManager):
        self.storage = storage_manager
        self.parsers = [
            PDFParser(),
            DOCXParser(),
            TextParser()
        ]
    
    def _get_parser(self, filename: str) -> Optional[DocumentParser]:
        """Get appropriate parser for file"""
        for parser in self.parsers:
            if parser.can_parse(filename):
                return parser
        return None
    
    def _extract_faqs_from_text(self, text: str) -> List[Tuple[str, str]]:
        """Extract FAQ pairs from text using simple heuristics"""
        import re
        
        faqs = []
        
        # Pattern 1: Q: ... A: ...
        qa_pattern = r'Q:\s*(.+?)\s*A:\s*(.+?)(?=Q:|$)'
        matches = re.findall(qa_pattern, text, re.DOTALL | re.IGNORECASE)
        for question, answer in matches:
            faqs.append((question.strip(), answer.strip()))
        
        # Pattern 2: Question: ... Answer: ...
        qa_pattern2 = r'Question:\s*(.+?)\s*Answer:\s*(.+?)(?=Question:|$)'
        matches2 = re.findall(qa_pattern2, text, re.DOTALL | re.IGNORECASE)
        for question, answer in matches2:
            faqs.append((question.strip(), answer.strip()))
        
        # Pattern 3: Lines ending with ? followed by lines
        lines = text.split('\n')
        current_question = None
        current_answer = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.endswith('?'):
                # Save previous FAQ if exists
                if current_question and current_answer:
                    faqs.append((current_question, ' '.join(current_answer)))
                
                current_question = line
                current_answer = []
            
            elif current_question and not line.endswith('?'):
                current_answer.append(line)
        
        # Save final FAQ
        if current_question and current_answer:
            faqs.append((current_question, ' '.join(current_answer)))
        
        return faqs
    
    def ingest_documents(
        self,
        project_id: str,
        files: List[Tuple[str, bytes]],  # (filename, content)
        ingest_type: str = "kb",
        split_mode: str = "tokens",
        max_chunk_chars: int = 1200
    ) -> Tuple[List[str], List[str], str]:
        """
        Ingest documents into project
        Returns: (created_ids, updated_ids, job_id)
        """
        job_id = str(uuid.uuid4())
        created_ids = []
        updated_ids = []
        
        for filename, content in files:
            try:
                # Save attachment
                self.storage.save_attachment(project_id, filename, content)
                
                # Parse document
                parser = self._get_parser(filename)
                if not parser:
                    continue  # Skip unsupported files
                
                text = parser.parse(content, filename)
                if not text.strip():
                    continue  # Skip empty documents
                
                if ingest_type == "faq":
                    # Try to extract FAQs
                    faq_pairs = self._extract_faqs_from_text(text)
                    
                    if faq_pairs:
                        # Create FAQ entries
                        faqs = []
                        for question, answer in faq_pairs:
                            faq = FAQEntry.from_qa(
                                project_id=project_id,
                                question=question,
                                answer=answer,
                                source="upload",
                                source_file=filename
                            )
                            faqs.append(faq)
                        
                        # Upsert FAQs
                        c_ids, u_ids = self.storage.upsert_faqs(project_id, faqs)
                        created_ids.extend(c_ids)
                        updated_ids.extend(u_ids)
                    
                    else:
                        # If no FAQs found, treat as KB
                        ingest_type = "kb"
                
                if ingest_type == "kb":
                    # Process as knowledge base content
                    if split_mode == "headings":
                        sections = TextSplitter.split_by_headings(text)
                        kb_entries = []
                        
                        for heading, content in sections:
                            if content.strip():
                                entry = KBEntry.from_content(
                                    project_id=project_id,
                                    article=f"{Path(filename).stem} - {heading}",
                                    content=content,
                                    source="upload",
                                    source_file=filename
                                )
                                kb_entries.append(entry)
                    
                    elif split_mode == "tokens":
                        chunks = TextSplitter.split_by_tokens(text, max_chunk_chars)
                        kb_entries = []
                        
                        for i, chunk in enumerate(chunks):
                            if chunk.strip():
                                article_name = Path(filename).stem
                                if len(chunks) > 1:
                                    article_name += f" - Part {i+1}"
                                
                                entry = KBEntry.from_content(
                                    project_id=project_id,
                                    article=article_name,
                                    content=chunk,
                                    source="upload",
                                    source_file=filename,
                                    chunk_index=i if len(chunks) > 1 else None
                                )
                                kb_entries.append(entry)
                    
                    else:  # No splitting
                        kb_entries = [KBEntry.from_content(
                            project_id=project_id,
                            article=Path(filename).stem,
                            content=text,
                            source="upload",
                            source_file=filename
                        )]
                    
                    # Upsert KB entries
                    c_ids, u_ids = self.storage.upsert_kb_entries(project_id, kb_entries)
                    created_ids.extend(c_ids)
                    updated_ids.extend(u_ids)
            
            except Exception as e:
                # Log error but continue with other files
                print(f"Error processing {filename}: {e}")
                continue
        
        return created_ids, updated_ids, job_id